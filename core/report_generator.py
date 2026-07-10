"""
GÖKTÜRK — Uyumluluk Raporu Üretici
Bulguları ve Annex 5 kapsam verisini, ISO/SAE 21434 / UN R155 referanslı
bir Word (.docx) raporuna dönüştürür. Rapor, çalışma anında (Streamlit
"Raporlama" sayfasından) gerçek FindingStore verisiyle üretilir.
"""

from datetime import datetime, timezone
from io import BytesIO
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from taxonomy.loader import load_taxonomy

IMPACT_ORDER = {"critical": 4, "high": 3, "medium": 2, "low": 1, "none": 0, "unknown": 0}
IMPACT_LABEL_TR = {
    "critical": "Kritik", "high": "Yüksek", "medium": "Orta",
    "low": "Düşük", "none": "Yok", "unknown": "Bilinmiyor",
}
STATUS_LABEL_TR = {
    "vulnerable": "Zafiyetli", "not_vulnerable": "Temiz",
    "inconclusive": "Sonuçsuz", "error": "Hata",
}


def _highest_impact(finding: Dict[str, Any]) -> str:
    impacts = [
        finding.get("impact_safety", "none"),
        finding.get("impact_financial", "none"),
        finding.get("impact_operational", "none"),
        finding.get("impact_privacy", "none"),
    ]
    return max(impacts, key=lambda x: IMPACT_ORDER.get(x, 0))


def _set_cell_text(cell, text: str, bold: bool = False, size: int = 10) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def generate_compliance_report(
    profile: Dict[str, Any],
    findings: List[Dict[str, Any]],
    coverage: Dict[str, Any],
) -> bytes:
    """Verilen profil/bulgu/kapsam verisinden bir docx uyumluluk raporu üretir.

    Args:
        profile: parsed vehicle profile dict (id, name, architecture, ...)
        findings: FindingStore.get_findings() çıktısı (dict listesi)
        coverage: FindingStore.get_compliance_coverage() çıktısı

    Returns:
        Raporun ham docx bayt içeriği (Streamlit st.download_button ile kullanılabilir).
    """
    doc = Document()

    # ── Kapak ────────────────────────────────────────────────────────────────
    title = doc.add_heading("GÖKTÜRK-AV — Siber Güvenlik Uyumluluk Raporu", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(profile.get("name", profile.get("id", "Bilinmeyen Araç")))
    run.bold = True
    run.font.size = Pt(16)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"Profil ID: {profile.get('id', '—')}  |  "
        f"Mimari: {profile.get('architecture', '—')}  |  "
        f"Üretim tarihi: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    ref = doc.add_paragraph()
    ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ref_run = ref.add_run("UN R155 Annex 5 (69 Saldırı Vektörü) ve ISO/SAE 21434 Clause 15 (TARA) referanslıdır")
    ref_run.italic = True
    ref_run.font.size = Pt(10)

    doc.add_page_break()

    # ── Yönetici Özeti ───────────────────────────────────────────────────────
    _add_heading(doc, "1. Yönetici Özeti", level=1)
    total = coverage.get("total", 0)
    vuln = coverage.get("vulnerable", 0)
    clean = coverage.get("not_vulnerable", 0)
    vectors_tested = coverage.get("vectors_tested", 0)

    summary_table = doc.add_table(rows=1, cols=4)
    summary_table.style = "Light Grid Accent 1"
    hdr = summary_table.rows[0].cells
    for i, label in enumerate(["Toplam Test", "Zafiyetli Bulgu", "Temiz Bulgu", "R155 Kapsamı"]):
        _set_cell_text(hdr[i], label, bold=True)
    row = summary_table.add_row().cells
    _set_cell_text(row[0], str(total))
    _set_cell_text(row[1], str(vuln))
    _set_cell_text(row[2], str(clean))
    _set_cell_text(row[3], f"{vectors_tested} / 69 (%{round(vectors_tested / 69 * 100, 1)})")

    doc.add_paragraph()
    critical_safety = [
        f for f in findings
        if f.get("status") == "vulnerable" and _highest_impact(f) in ("critical", "high")
    ]
    if critical_safety:
        warn = doc.add_paragraph()
        warn_run = warn.add_run(
            f"⚠ {len(critical_safety)} bulgu yüksek/kritik safety etkisi taşımaktadır "
            "ve öncelikli olarak ele alınmalıdır (bkz. Bölüm 3)."
        )
        warn_run.bold = True
        warn_run.font.color.rgb = RGBColor(0xB0, 0x1E, 0x1E)

    # ── Metodoloji ───────────────────────────────────────────────────────────
    _add_heading(doc, "2. Metodoloji", level=1)
    doc.add_paragraph(
        "Bu rapor, GÖKTÜRK-AV platformu tarafından UN R155 Annex 5 taksonomisine "
        "(69 saldırı vektörü, 7 tehdit kategorisi) çapalı olarak otomatik "
        "üretilmiştir. Her bulgu, ISO/SAE 21434 Clause 15 (TARA) kapsamında "
        "dört etki boyutunda (safety, financial, operational, privacy) ve "
        "Annex H saldırı fizibilitesi ölçeğinde (low/medium/high/very_high) "
        "değerlendirilmiştir."
    )
    doc.add_paragraph(
        "Bu rapor otomatik bir test aracının çıktısıdır; ISO 21434 uyumluluğu "
        "için gereken insan gözetimli TARA sürecinin yerini tutmaz. Bulgular, "
        "bir güvenlik mühendisi tarafından gözden geçirilmeli ve risk işleme "
        "kararları buna göre alınmalıdır."
    )

    # ── Bulgular ──────────────────────────────────────────────────────────────
    _add_heading(doc, "3. Bulgular", level=1)
    if not findings:
        doc.add_paragraph("Bu profil için henüz kaydedilmiş bulgu bulunmamaktadır.")
    else:
        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Grid Accent 1"
        headers = ["R155 Vektör", "Başlık", "Bileşen", "Durum", "En Yüksek Etki", "Fizibilite"]
        for i, h in enumerate(headers):
            _set_cell_text(table.rows[0].cells[i], h, bold=True)

        for f in sorted(findings, key=lambda x: (x.get("r155_vector_id") or "", x.get("component_id") or "")):
            cells = table.add_row().cells
            _set_cell_text(cells[0], f.get("r155_vector_id") or "—")
            _set_cell_text(cells[1], f.get("title", ""))
            _set_cell_text(cells[2], f.get("component_id", ""))
            _set_cell_text(cells[3], STATUS_LABEL_TR.get(f.get("status"), f.get("status", "—")))
            _set_cell_text(cells[4], IMPACT_LABEL_TR.get(_highest_impact(f), "—"))
            _set_cell_text(cells[5], f.get("attack_feasibility") or "—")

        # Kritik/yüksek zafiyetli bulgular için detaylı alt bölüm
        # Aynı bulgu birden fazla bileşende tekrarlanabilir (ör. LiDAR spoof
        # dört sensörde de çalışır) — metni tekrar tekrar basmak yerine
        # (vektör, başlık, açıklama, remediation) bazında gruplayıp etkilenen
        # bileşenleri tek satırda listeliyoruz.
        if critical_safety:
            doc.add_paragraph()
            _add_heading(doc, "3.1 Öncelikli Bulgu Detayları (Yüksek/Kritik Safety)", level=2)

            grouped: Dict[tuple, Dict[str, Any]] = {}
            order: List[tuple] = []
            for f in critical_safety:
                key = (f.get("r155_vector_id", ""), f.get("title", ""), f.get("description", ""))
                if key not in grouped:
                    grouped[key] = {
                        "components": [],
                        "remediation": f.get("remediation", ""),
                    }
                    order.append(key)
                grouped[key]["components"].append(f.get("component_id", "—"))

            for key in order:
                vector_id, title_text, description = key
                info = grouped[key]
                p = doc.add_paragraph()
                r = p.add_run(f"{vector_id} — {title_text}")
                r.bold = True
                comp_p = doc.add_paragraph()
                comp_label = comp_p.add_run("Etkilenen bileşenler: ")
                comp_label.bold = True
                comp_p.add_run(", ".join(info["components"]))
                doc.add_paragraph(description)
                if info["remediation"]:
                    rem = doc.add_paragraph()
                    rem_label = rem.add_run("Önerilen İyileştirme: ")
                    rem_label.bold = True
                    rem.add_run(info["remediation"])
                doc.add_paragraph()

    # ── Annex 5 Kapsam ────────────────────────────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, "4. UN R155 Annex 5 Kapsam Durumu", level=1)
    taxonomy = load_taxonomy()
    tested_vectors = {f.get("r155_vector_id") for f in findings if f.get("r155_vector_id")}

    cov_table = doc.add_table(rows=1, cols=4)
    cov_table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Kategori", "Kategori Adı", "Test Edilen / Toplam Vektör", "Kapsam"]):
        _set_cell_text(cov_table.rows[0].cells[i], h, bold=True)

    for cat in taxonomy["categories"]:
        cat_vector_ids = {v["id"] for v in cat["vectors"]}
        cat_tested = cat_vector_ids & tested_vectors
        cells = cov_table.add_row().cells
        _set_cell_text(cells[0], cat["code"])
        _set_cell_text(cells[1], cat["label"])
        _set_cell_text(cells[2], f"{len(cat_tested)} / {len(cat_vector_ids)}")
        pct = round(len(cat_tested) / len(cat_vector_ids) * 100) if cat_vector_ids else 0
        _set_cell_text(cells[3], f"%{pct}")

    doc.add_paragraph()
    total_pct = round(len(tested_vectors) / taxonomy["total_vectors"] * 100, 1)
    footer_p = doc.add_paragraph()
    footer_run = footer_p.add_run(
        f"Toplam Annex 5 kapsamı: {len(tested_vectors)} / {taxonomy['total_vectors']} "
        f"vektör (%{total_pct})"
    )
    footer_run.bold = True

    # ── Alt bilgi / feragatname ──────────────────────────────────────────────
    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disc_run = disclaimer.add_run(
        "Bu rapor GÖKTÜRK-AV tarafından otomatik üretilmiştir. Tip onayı veya "
        "resmi uyumluluk beyanı için bir güvenlik mühendisi tarafından "
        "doğrulanması gerekir."
    )
    disc_run.italic = True
    disc_run.font.size = Pt(9)
    disc_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
