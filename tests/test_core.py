"""
GÖKTÜRK — Çekirdek Test Suite (Faz 2-4)
Tüm katmanları kapsar: taksonomi, Finding şeması, FindingStore,
mock adaptör, yedi test modülü, orchestrator, rapor üretici ve
3D saldırı yüzeyi haritası.

Donanımsızdır; her şey MockAdapter üzerinden deterministik koşar.
"""

import pytest

from taxonomy.loader import load_taxonomy, get_vector, get_category
from plugins.base_plugin import Finding, BasePlugin
from core.finding_store import FindingStore
from core.orchestrator import Orchestrator
from adapters.mock_adapter import MockAdapter

from plugins.modules.can_replay_plugin import CANReplayPlugin
from plugins.modules.can_fuzz_plugin import CANFuzzPlugin
from plugins.modules.ros2_topic_enum_plugin import ROS2TopicEnumPlugin
from plugins.modules.gps_spoof_plugin import GPSSpoofPlugin
from plugins.modules.obd2_enum_plugin import OBD2EnumPlugin
from plugins.modules.ros2_topic_injection_plugin import ROS2TopicInjectionPlugin
from plugins.modules.lidar_spoof_plugin import LidarSpoofPlugin
from plugins.modules.v2x_spoof_plugin import V2XSpoofPlugin
from plugins.modules.ecu_fuzz_plugin import ECUFuzzPlugin
from plugins.modules.ota_attack_plugin import OTAAttackPlugin
from plugins.modules.adversarial_ml_plugin import AdversarialMLPlugin
from plugins.modules.backend_server_plugin import BackendServerPlugin
from plugins.modules.diag_access_abuse_plugin import DiagnosticAccessAbusePlugin
from plugins.modules.debug_port_access_plugin import DebugPortAccessPlugin
from plugins.modules.firmware_integrity_plugin import FirmwareIntegrityPlugin
from plugins.modules.remote_telematics_exploit_plugin import RemoteTelematicsExploitPlugin
from plugins.modules.can_dos_plugin import CANDosPlugin
from plugins.modules.ivi_pivot_plugin import IVIPivotPlugin
from plugins.modules.telematics_channel_plugin import TelematicsChannelPlugin
from plugins.modules.physical_ecu_access_plugin import PhysicalECUAccessPlugin
from core.report_generator import generate_compliance_report
from core.attack_surface import compute_component_statuses, build_attack_surface_html
from core.compliance_heatmap import compute_vector_statuses, build_heatmap_html
from docx import Document
from io import BytesIO


# ── Fixtures ─────────────────────────────────────────────────────────────────

@pytest.fixture
def profile():
    return {
        "id": "shuttle-test",
        "name": "Test Shuttle",
        "components": [
            {"id": "gateway_ecu", "attack_surfaces": ["in-vehicle-network"],
             "r155_vectors": ["R155-2.2", "R155-2.5"]},
            {"id": "hpc", "attack_surfaces": ["ros2-dds"], "r155_vectors": ["R155-5.6"]},
            {"id": "gnss", "attack_surfaces": ["sensor-spoofing"], "r155_vectors": ["R155-2.8"]},
            {"id": "obd2_port", "attack_surfaces": ["diagnostic"], "r155_vectors": ["R155-5.5"]},
        ],
    }


def _mock(mode="vulnerable"):
    a = MockAdapter({"mode": mode})
    a.connect()
    return a


# ── Taksonomi ────────────────────────────────────────────────────────────────

def test_taxonomy_loads():
    tax = load_taxonomy()
    assert "categories" in tax
    assert len(tax["categories"]) == 7


def test_taxonomy_total_vectors_is_69():
    tax = load_taxonomy()
    total = sum(len(cat["vectors"]) for cat in tax["categories"])
    assert total == 69
    assert tax["total_vectors"] == 69


def test_taxonomy_vector_ids_unique():
    tax = load_taxonomy()
    ids = [v["id"] for cat in tax["categories"] for v in cat["vectors"]]
    assert len(ids) == len(set(ids))


def test_get_vector_returns_category_context():
    v = get_vector("R155-2.5")
    assert v["id"] == "R155-2.5"
    assert v["category_id"] == 2
    assert "category_label" in v


def test_get_vector_unknown_returns_empty():
    assert get_vector("R155-9.9") == {}


def test_get_category_returns_dict():
    cat = get_category(5)
    assert cat["id"] == 5
    assert len(cat["vectors"]) > 0


def test_get_category_unknown_returns_empty():
    assert get_category(99) == {}


# ── Finding şeması ───────────────────────────────────────────────────────────

def test_finding_defaults():
    f = Finding(component_id="c", test_module_id="m", status="vulnerable", title="t")
    assert f.impact_safety == "none"
    assert f.attack_feasibility == "unknown"
    assert f.id and f.created_at


def test_finding_is_vulnerable():
    f = Finding(component_id="c", test_module_id="m", status="vulnerable", title="t")
    assert f.is_vulnerable() is True


def test_finding_is_critical_safety():
    f = Finding(component_id="c", test_module_id="m", status="vulnerable",
                title="t", impact_safety="high")
    assert f.is_critical_safety() is True


def test_finding_not_critical_when_clean():
    f = Finding(component_id="c", test_module_id="m", status="not_vulnerable",
                title="t", impact_safety="high")
    assert f.is_critical_safety() is False


def test_finding_highest_impact():
    f = Finding(component_id="c", test_module_id="m", status="vulnerable", title="t",
                impact_safety="medium", impact_operational="critical")
    assert f.highest_impact() == "critical"


def test_finding_to_dict_roundtrip():
    f = Finding(component_id="c", test_module_id="m", status="vulnerable", title="t")
    d = f.to_dict()
    assert d["component_id"] == "c"
    assert d["status"] == "vulnerable"


# ── FindingStore ─────────────────────────────────────────────────────────────

def test_store_init(tmp_path):
    db = FindingStore(db_path=str(tmp_path / "t.db"))
    assert db is not None


def test_store_profile_roundtrip(tmp_path):
    db = FindingStore(db_path=str(tmp_path / "t.db"))
    db.save_profile("p1", "Prof 1", "id: p1")
    assert db.get_profile("p1")["name"] == "Prof 1"
    assert any(p["id"] == "p1" for p in db.get_profiles())


def test_store_session_lifecycle(tmp_path):
    db = FindingStore(db_path=str(tmp_path / "t.db"))
    db.save_profile("p1", "Prof 1", "x")
    sid = db.create_session("p1", "not")
    db.close_session(sid)
    sessions = db.get_sessions("p1")
    assert sessions[0]["status"] == "completed"


def test_store_add_and_get_finding(tmp_path):
    db = FindingStore(db_path=str(tmp_path / "t.db"))
    db.save_profile("p1", "Prof 1", "x")
    sid = db.create_session("p1")
    db.add_finding(sid, "p1", "comp", "mod", "vulnerable", "T",
                   r155_vector_id="R155-2.5", r155_category=2)
    findings = db.get_findings(vehicle_profile_id="p1")
    assert len(findings) == 1
    assert findings[0]["r155_vector_id"] == "R155-2.5"
    assert findings[0]["evidence_paths"] == []


def test_store_filter_by_status(tmp_path):
    db = FindingStore(db_path=str(tmp_path / "t.db"))
    db.save_profile("p1", "Prof 1", "x")
    sid = db.create_session("p1")
    db.add_finding(sid, "p1", "c", "m", "vulnerable", "A")
    db.add_finding(sid, "p1", "c", "m", "not_vulnerable", "B")
    assert len(db.get_findings(status="vulnerable")) == 1


def test_store_compliance_coverage(tmp_path):
    db = FindingStore(db_path=str(tmp_path / "t.db"))
    db.save_profile("p1", "Prof 1", "x")
    sid = db.create_session("p1")
    db.add_finding(sid, "p1", "c", "m", "vulnerable", "A", r155_vector_id="R155-2.5")
    db.add_finding(sid, "p1", "c", "m", "not_vulnerable", "B", r155_vector_id="R155-5.6")
    cov = db.get_compliance_coverage("p1")
    assert cov["total"] == 2
    assert cov["vulnerable"] == 1
    assert cov["vectors_tested"] == 2


def test_store_module_registry(tmp_path):
    db = FindingStore(db_path=str(tmp_path / "t.db"))
    db.register_module({"id": "can-replay", "name": "Replay",
                        "surface": "in-vehicle-network", "technique": "replay"})
    mods = db.get_modules()
    assert any(m["id"] == "can-replay" for m in mods)


# ── MockAdapter ──────────────────────────────────────────────────────────────

def test_mock_invalid_mode_raises():
    with pytest.raises(ValueError):
        MockAdapter({"mode": "banana"})


def test_mock_connect_and_impersonate():
    a = MockAdapter({"mode": "secure", "as_type": "socketcan"})
    assert a.connect() is True
    assert a.is_connected() is True
    assert a.adapter_type == "socketcan"


def test_mock_modes_behaviour():
    assert _mock("vulnerable").receive_frames() != []
    assert _mock("empty").receive_frames() == []
    assert _mock("vulnerable").send_frame({"arb_id": 1}) is True
    assert _mock("secure").send_frame({"arb_id": 1}) is False
    assert _mock("secure").list_topics() and _mock("empty").list_topics() == []


def test_mock_v2x_injection_behaviour():
    # imzasız mesaj: sadece vulnerable kabul; imzalı: her modda kabul
    assert _mock("vulnerable").inject_v2x_message(signed=False) is True
    assert _mock("secure").inject_v2x_message(signed=False) is False
    assert _mock("empty").inject_v2x_message(signed=False) is False
    assert _mock("secure").inject_v2x_message(signed=True) is True


def test_mock_v2x_attack_probe_behaviour():
    # vulnerable -> her senaryo accepted; secure -> hicbiri; empty -> yanitsiz
    for scn in ("identity_spoof", "v2i_infra_trust"):
        assert _mock("vulnerable").v2x_attack_probe("v2x_obu", scn)["accepted"] is True
        assert _mock("secure").v2x_attack_probe("v2x_obu", scn)["accepted"] is False
        assert _mock("empty").v2x_attack_probe("v2x_obu", scn)["accepted"] is False


def test_mock_fuzz_ecu_behaviour():
    # vulnerable -> girdiler islenir + fault uretir; secure -> hepsi reddedilir; empty -> bos
    vuln = _mock("vulnerable").fuzz_ecu("hpc", mode="smart", count=100)
    assert vuln and any(r["memory_fault"] for r in vuln)
    assert all(r["accepted"] for r in vuln)
    sec = _mock("secure").fuzz_ecu("hpc", count=100)
    assert sec and not any(r["accepted"] for r in sec)
    assert _mock("empty").fuzz_ecu("hpc") == []


def test_mock_ota_probe_behaviour():
    # vulnerable -> her senaryo accepted; secure -> hicbiri; empty -> yanitsiz
    for scn in ("rollback", "bad_signature", "plaintext", "pre_update_tamper", "manifest_tamper"):
        assert _mock("vulnerable").ota_update_probe("tcu", scn)["accepted"] is True
        assert _mock("secure").ota_update_probe("tcu", scn)["accepted"] is False
        assert _mock("empty").ota_update_probe("tcu", scn)["accepted"] is False


def test_mock_adversarial_perturbation_behaviour():
    # vulnerable -> fooled; secure -> defended (fooled degil); empty -> yanit yok
    v = _mock("vulnerable").inject_adversarial_perturbation("cam", sensor="camera")
    assert v["fooled"] is True and v["defended"] is False
    assert v["original"] != v["adversarial"]
    s = _mock("secure").inject_adversarial_perturbation("cam", sensor="camera")
    assert s["fooled"] is False and s["defended"] is True
    assert s["original"] == s["adversarial"]
    e = _mock("empty").inject_adversarial_perturbation("cam", sensor="camera")
    assert e["fooled"] is False and e["defended"] is False


def test_mock_backend_server_probe_behaviour():
    for scn in ("weak_auth", "dos", "supply_chain_compromise"):
        assert _mock("vulnerable").backend_server_probe("telematics_module", scn)["accepted"] is True
        assert _mock("secure").backend_server_probe("telematics_module", scn)["accepted"] is False
        assert _mock("empty").backend_server_probe("telematics_module", scn)["accepted"] is False


def test_mock_diagnostic_scope_probe_behaviour():
    assert _mock("vulnerable").diagnostic_scope_probe("obd2_port") is True
    assert _mock("secure").diagnostic_scope_probe("obd2_port") is False
    assert _mock("empty").diagnostic_scope_probe("obd2_port") is False


def test_mock_debug_port_probe_behaviour():
    assert _mock("vulnerable").debug_port_probe("debug_ports") is True
    assert _mock("secure").debug_port_probe("debug_ports") is False
    assert _mock("empty").debug_port_probe("debug_ports") is False


def test_mock_physical_ecu_access_probe_behaviour():
    for method in ("enclosure_bypass", "harness_tap"):
        assert _mock("vulnerable").physical_ecu_access_probe("obd2_port", method=method) is True
        assert _mock("secure").physical_ecu_access_probe("obd2_port", method=method) is False
        assert _mock("empty").physical_ecu_access_probe("obd2_port", method=method) is False


def test_mock_firmware_integrity_probe_behaviour():
    for scn in ("malicious_replace", "integrity_check_bypass", "secure_boot_bypass"):
        assert _mock("vulnerable").firmware_integrity_probe("hpc_compute", scn)["accepted"] is True
        assert _mock("secure").firmware_integrity_probe("hpc_compute", scn)["accepted"] is False
        assert _mock("empty").firmware_integrity_probe("hpc_compute", scn)["accepted"] is False


# ── Plugin katmanı ───────────────────────────────────────────────────────────

def test_can_replay_matrix():
    assert CANReplayPlugin(_mock("vulnerable")).run({"id": "c"}).status == "vulnerable"
    assert CANReplayPlugin(_mock("secure")).run({"id": "c"}).status == "not_vulnerable"
    assert CANReplayPlugin(_mock("empty")).run({"id": "c"}).status == "inconclusive"


def test_can_fuzz_matrix():
    assert CANFuzzPlugin(_mock("vulnerable")).run({"id": "c"}).status == "vulnerable"
    assert CANFuzzPlugin(_mock("secure")).run({"id": "c"}).status == "not_vulnerable"
    assert CANFuzzPlugin(_mock("empty")).run({"id": "c"}).status == "inconclusive"


def test_ros2_enum_matrix():
    assert ROS2TopicEnumPlugin(_mock("vulnerable")).run({"id": "c"}).status == "vulnerable"
    assert ROS2TopicEnumPlugin(_mock("secure")).run({"id": "c"}).status == "not_vulnerable"
    assert ROS2TopicEnumPlugin(_mock("empty")).run({"id": "c"}).status == "inconclusive"


def test_gps_spoof_matrix():
    assert GPSSpoofPlugin(_mock("vulnerable")).run({"id": "c"}).status == "vulnerable"
    assert GPSSpoofPlugin(_mock("secure")).run({"id": "c"}).status == "not_vulnerable"


def test_obd2_enum_matrix():
    assert OBD2EnumPlugin(_mock("vulnerable")).run({"id": "c"}).status == "vulnerable"
    assert OBD2EnumPlugin(_mock("secure")).run({"id": "c"}).status == "not_vulnerable"
    assert OBD2EnumPlugin(_mock("empty")).run({"id": "c"}).status == "inconclusive"


def test_vulnerable_findings_carry_taxonomy():
    f = OBD2EnumPlugin(_mock("vulnerable")).run({"id": "c"})
    assert f.r155_vector_id == "R155-5.5"
    assert f.r155_category == 5
    assert f.is_vulnerable()


def test_plugin_never_raises_on_bad_component():
    # component_config eksik olsa da plugin exception fırlatmamalı
    f = CANReplayPlugin(_mock("vulnerable")).run({})
    assert isinstance(f, Finding)


def test_ros2_injection_matrix():
    assert ROS2TopicInjectionPlugin(_mock("vulnerable")).run({"id": "c"}).status == "vulnerable"
    assert ROS2TopicInjectionPlugin(_mock("secure")).run({"id": "c"}).status == "not_vulnerable"
    assert ROS2TopicInjectionPlugin(_mock("empty")).run({"id": "c"}).status == "inconclusive"


def test_ros2_injection_critical_safety_when_vulnerable():
    f = ROS2TopicInjectionPlugin(_mock("vulnerable")).run({"id": "c"})
    assert f.r155_vector_id == "R155-5.7"
    assert f.impact_safety == "critical"
    assert f.is_critical_safety()


def test_lidar_spoof_matrix():
    assert LidarSpoofPlugin(_mock("vulnerable")).run({"id": "lidar_front"}).status == "vulnerable"
    assert LidarSpoofPlugin(_mock("secure")).run({"id": "lidar_front"}).status == "not_vulnerable"


def test_lidar_spoof_empty_mode_rejected():
    # empty modda inject_lidar_spoof her iki senaryoda da False döner → not_vulnerable
    f = LidarSpoofPlugin(_mock("empty")).run({"id": "lidar_front"})
    assert f.status == "not_vulnerable"


def test_lidar_spoof_remove_scenario_is_critical():
    f = LidarSpoofPlugin(_mock("vulnerable")).run({"id": "lidar_front"})
    assert f.impact_safety == "critical"
    assert f.is_critical_safety()


def test_v2x_spoof_matrix():
    vuln = V2XSpoofPlugin(_mock("vulnerable")).run({"id": "v2x_unit"})
    assert isinstance(vuln, list) and all(f.status == "vulnerable" for f in vuln)
    assert V2XSpoofPlugin(_mock("secure")).run({"id": "v2x_unit"}).status == "not_vulnerable"
    assert V2XSpoofPlugin(_mock("empty")).run({"id": "v2x_unit"}).status == "not_vulnerable"


def test_v2x_spoof_vulnerable_returns_three_distinct_vectors():
    findings = V2XSpoofPlugin(_mock("vulnerable")).run({"id": "v2x_obu"})
    vectors = sorted(f.r155_vector_id for f in findings)
    assert vectors == ["R155-2.1", "R155-2.7", "R155-5.12"]
    assert all(f.r155_category == 2 for f in findings)
    assert all(f.is_vulnerable() for f in findings)
    bsm_finding = next(f for f in findings if f.r155_vector_id == "R155-2.7")
    assert bsm_finding.impact_safety == "high"


def test_v2x_spoof_secure_reports_all_protected():
    f = V2XSpoofPlugin(_mock("secure")).run({"id": "v2x_obu"})
    assert f.status == "not_vulnerable"
    assert "korumalı" in f.description.lower() or "reddedildi" in f.title.lower()


def test_v2x_spoof_inconclusive_when_adapter_unsupported():
    # Ham base adapter inject_v2x_message'i desteklemez -> NotImplementedError -> inconclusive
    from adapters.base_adapter import BaseAdapter

    class _BareAdapter(BaseAdapter):
        adapter_type = "bare"

        def connect(self):
            self._connected = True
            return True

        def disconnect(self):
            self._connected = False

        def is_connected(self):
            return self._connected

    bare = _BareAdapter({})
    bare.connect()
    f = V2XSpoofPlugin(bare).run({"id": "v2x_unit"})
    assert f.status == "inconclusive"


def test_ecu_fuzz_matrix():
    vuln = ECUFuzzPlugin(_mock("vulnerable")).run({"id": "hpc"})
    assert isinstance(vuln, list) and all(f.status == "vulnerable" for f in vuln)
    assert ECUFuzzPlugin(_mock("secure")).run({"id": "hpc"}).status == "not_vulnerable"
    assert ECUFuzzPlugin(_mock("empty")).run({"id": "hpc"}).status == "inconclusive"


def test_ecu_fuzz_splits_memory_fault_and_hang_vectors():
    findings = ECUFuzzPlugin(_mock("vulnerable")).run({"id": "hpc"})
    vectors = sorted(f.r155_vector_id for f in findings)
    assert vectors == ["R155-6.8", "R155-6.9"]
    assert all(f.r155_category == 6 for f in findings)
    assert all(f.is_vulnerable() for f in findings)
    mem_finding = next(f for f in findings if f.r155_vector_id == "R155-6.8")
    assert mem_finding.impact_safety == "high"
    hang_finding = next(f for f in findings if f.r155_vector_id == "R155-6.9")
    assert hang_finding.impact_safety == "high"
    assert "hang" in hang_finding.title.lower() or "yanıtsız" in hang_finding.title.lower()


def test_ecu_fuzz_reports_memory_fault_in_title():
    findings = ECUFuzzPlugin(_mock("vulnerable")).run({"id": "hpc"})
    mem_finding = next(f for f in findings if f.r155_vector_id == "R155-6.8")
    assert "fault" in mem_finding.title.lower() or "bozulma" in mem_finding.title.lower()


def test_ecu_fuzz_mode_config_respected():
    plugin = ECUFuzzPlugin(_mock("vulnerable"), config={"fuzz_mode": "dumb", "count": 200})
    findings = plugin.run({"id": "hpc"})
    assert isinstance(findings, list) and all(f.status == "vulnerable" for f in findings)


def test_ota_attack_matrix():
    vuln = OTAAttackPlugin(_mock("vulnerable")).run({"id": "tcu"})
    assert isinstance(vuln, list) and all(f.status == "vulnerable" for f in vuln)
    assert OTAAttackPlugin(_mock("secure")).run({"id": "tcu"}).status == "not_vulnerable"
    assert OTAAttackPlugin(_mock("empty")).run({"id": "tcu"}).status == "not_vulnerable"


def test_ota_attack_vulnerable_returns_five_distinct_vectors():
    findings = OTAAttackPlugin(_mock("vulnerable")).run({"id": "tcu"})
    vectors = sorted(f.r155_vector_id for f in findings)
    assert vectors == ["R155-3.1", "R155-3.4", "R155-3.5", "R155-3.6", "R155-3.7"]
    assert all(f.r155_category == 3 for f in findings)
    sig_finding = next(f for f in findings if f.r155_vector_id == "R155-3.4")
    assert sig_finding.impact_safety == "high"
    pre_update_finding = next(f for f in findings if f.r155_vector_id == "R155-3.1")
    assert pre_update_finding.impact_safety == "high"
    manifest_finding = next(f for f in findings if f.r155_vector_id == "R155-3.7")
    assert manifest_finding.impact_safety == "medium"


def test_ota_attack_lists_all_five_scenarios():
    findings = OTAAttackPlugin(_mock("vulnerable")).run({"id": "tcu"})
    assert len(findings) == 5


def test_ota_attack_secure_reports_all_protected():
    f = OTAAttackPlugin(_mock("secure")).run({"id": "tcu"})
    assert f.status == "not_vulnerable"
    assert "korumalar" in f.title.lower() or "aktif" in f.title.lower()


def test_adversarial_ml_matrix():
    assert AdversarialMLPlugin(_mock("vulnerable")).run({"id": "camera_processor"}).status == "vulnerable"
    assert AdversarialMLPlugin(_mock("secure")).run({"id": "camera_processor"}).status == "not_vulnerable"
    assert AdversarialMLPlugin(_mock("empty")).run({"id": "camera_processor"}).status == "inconclusive"


def test_adversarial_ml_carries_taxonomy_and_critical_safety():
    f = AdversarialMLPlugin(_mock("vulnerable")).run({"id": "camera_processor"})
    assert f.r155_vector_id == "R155-6.7"
    assert f.r155_category == 6
    assert f.impact_safety == "critical"
    assert f.is_critical_safety()


def test_adversarial_ml_title_shows_misclassification():
    f = AdversarialMLPlugin(_mock("vulnerable")).run({"id": "camera_processor"})
    assert "\u2192" in f.title or "->" in f.title


def test_adversarial_ml_lidar_sensor_config():
    plugin = AdversarialMLPlugin(_mock("vulnerable"), config={"adv_sensor": "lidar"})
    f = plugin.run({"id": "lidar_controller"})
    assert f.status == "vulnerable"
    assert "lidar" in f.title.lower()


def test_backend_server_matrix():
    vuln = BackendServerPlugin(_mock("vulnerable")).run({"id": "telematics_module"})
    assert isinstance(vuln, list) and all(f.status == "vulnerable" for f in vuln)
    assert BackendServerPlugin(_mock("secure")).run({"id": "telematics_module"}).status == "not_vulnerable"
    assert BackendServerPlugin(_mock("empty")).run({"id": "telematics_module"}).status == "not_vulnerable"


def test_backend_server_vulnerable_returns_three_distinct_vectors():
    findings = BackendServerPlugin(_mock("vulnerable")).run({"id": "telematics_module"})
    vectors = sorted(f.r155_vector_id for f in findings)
    assert vectors == ["R155-1.1", "R155-1.4", "R155-1.5"]
    assert all(f.r155_category == 1 for f in findings)
    auth_finding = next(f for f in findings if f.r155_vector_id == "R155-1.1")
    assert auth_finding.impact_safety == "high"
    assert auth_finding.impact_privacy == "high"
    supply_chain_finding = next(f for f in findings if f.r155_vector_id == "R155-1.4")
    assert supply_chain_finding.impact_privacy == "high"


def test_backend_server_lists_all_three_scenarios():
    findings = BackendServerPlugin(_mock("vulnerable")).run({"id": "telematics_module"})
    assert len(findings) == 3


def test_backend_server_secure_reports_all_protected():
    f = BackendServerPlugin(_mock("secure")).run({"id": "telematics_module"})
    assert f.status == "not_vulnerable"
    assert "korumalar" in f.title.lower() or "aktif" in f.title.lower()


def test_diag_access_abuse_matrix():
    assert DiagnosticAccessAbusePlugin(_mock("vulnerable")).run({"id": "obd2_port"}).status == "vulnerable"
    assert DiagnosticAccessAbusePlugin(_mock("secure")).run({"id": "obd2_port"}).status == "not_vulnerable"
    assert DiagnosticAccessAbusePlugin(_mock("empty")).run({"id": "obd2_port"}).status == "not_vulnerable"


def test_diag_access_abuse_carries_taxonomy():
    f = DiagnosticAccessAbusePlugin(_mock("vulnerable")).run({"id": "obd2_port"})
    assert f.r155_vector_id == "R155-4.2"
    assert f.r155_category == 4
    assert f.is_vulnerable()


def test_diag_access_abuse_bulk_extract_config():
    plugin = DiagnosticAccessAbusePlugin(_mock("vulnerable"), config={"diag_action": "bulk_extract"})
    f = plugin.run({"id": "obd2_port"})
    assert f.status == "vulnerable"
    assert "veri cekme" in f.title.lower() or "veri çekme" in f.title.lower()


def test_debug_port_access_matrix():
    assert DebugPortAccessPlugin(_mock("vulnerable")).run({"id": "debug_ports"}).status == "vulnerable"
    assert DebugPortAccessPlugin(_mock("secure")).run({"id": "debug_ports"}).status == "not_vulnerable"
    assert DebugPortAccessPlugin(_mock("empty")).run({"id": "debug_ports"}).status == "not_vulnerable"


def test_debug_port_access_carries_taxonomy():
    f = DebugPortAccessPlugin(_mock("vulnerable")).run({"id": "debug_ports"})
    assert f.r155_vector_id == "R155-7.4"
    assert f.r155_category == 7
    assert f.is_vulnerable()
    assert f.attack_feasibility == "low"


def test_debug_port_access_uart_config():
    plugin = DebugPortAccessPlugin(_mock("vulnerable"), config={"debug_action": "uart_console"})
    f = plugin.run({"id": "debug_ports"})
    assert f.status == "vulnerable"
    assert "uart" in f.title.lower()


def test_physical_ecu_access_matrix():
    assert PhysicalECUAccessPlugin(_mock("vulnerable")).run({"id": "obd2_port"}).status == "vulnerable"
    assert PhysicalECUAccessPlugin(_mock("secure")).run({"id": "obd2_port"}).status == "not_vulnerable"
    assert PhysicalECUAccessPlugin(_mock("empty")).run({"id": "obd2_port"}).status == "not_vulnerable"


def test_physical_ecu_access_carries_taxonomy():
    f = PhysicalECUAccessPlugin(_mock("vulnerable")).run({"id": "obd2_port"})
    assert f.r155_vector_id == "R155-7.1"
    assert f.r155_category == 7
    assert f.is_vulnerable()
    assert f.attack_feasibility == "low"


def test_physical_ecu_access_harness_tap_config():
    plugin = PhysicalECUAccessPlugin(_mock("vulnerable"), config={"physical_access_method": "harness_tap"})
    f = plugin.run({"id": "obd2_port"})
    assert f.status == "vulnerable"
    assert "splice" in f.title.lower() or "tap" in f.title.lower()


def test_base_plugin_is_abstract():
    with pytest.raises(TypeError):
        BasePlugin(_mock())  # abstract run() → örneklenemez


# ── Orchestrator ─────────────────────────────────────────────────────────────

def test_orchestrator_discovers_all_plugins():
    orch = Orchestrator(_mock(), FindingStore(db_path=":memory:"), strict_adapter=False)
    classes = orch.discover_plugin_classes()
    ids = {c.module_id for c in classes}
    assert {"can-replay", "can-fuzz", "ros2-topic-enum", "ros2-topic-injection",
            "gps-spoof", "obd2-enum", "lidar-spoof", "v2x-spoof", "ecu-fuzz", "ota-attack",
            "adversarial-ml", "backend-server", "diag-access-abuse",
            "debug-port-access", "firmware-integrity", "physical-ecu-access"} <= ids


def test_orchestrator_run_persists_findings(tmp_path, profile):
    db = FindingStore(db_path=str(tmp_path / "run.db"))
    orch = Orchestrator(_mock("vulnerable"), db, strict_adapter=False)
    findings = orch.run_all(profile)
    assert len(findings) >= 7
    assert all(isinstance(f, Finding) for f in findings)
    stored = db.get_findings(vehicle_profile_id="shuttle-test")
    assert len(stored) == len(findings)


def test_orchestrator_vulnerable_profile_flags_risks(tmp_path, profile):
    db = FindingStore(db_path=str(tmp_path / "run.db"))
    orch = Orchestrator(_mock("vulnerable"), db, strict_adapter=False)
    orch.run_all(profile)
    cov = db.get_compliance_coverage("shuttle-test")
    assert cov["vulnerable"] >= 7


def test_orchestrator_secure_profile_no_vulns(tmp_path, profile):
    db = FindingStore(db_path=str(tmp_path / "run.db"))
    orch = Orchestrator(_mock("secure"), db, strict_adapter=False)
    orch.run_all(profile)
    cov = db.get_compliance_coverage("shuttle-test")
    assert cov["vulnerable"] == 0


def test_orchestrator_strict_adapter_skips_mismatch(tmp_path, profile):
    # strict + saf 'mock' tipi → hiçbir plugin'in applicable_adapters'ı 'mock' değil
    db = FindingStore(db_path=str(tmp_path / "run.db"))
    orch = Orchestrator(_mock("vulnerable"), db, strict_adapter=True)
    findings = orch.run_all(profile)
    assert all(f.status == "inconclusive" for f in findings)


# ── Rapor Üretici ─────────────────────────────────────────────────────────────

def _sample_finding(**overrides):
    base = {
        "component_id": "gateway_ecu",
        "test_module_id": "can-replay",
        "r155_vector_id": "R155-2.5",
        "r155_category": 2,
        "status": "vulnerable",
        "title": "CAN Replay: Kimlik doğrulama YOK",
        "description": "Açıklama metni.",
        "impact_safety": "high",
        "impact_financial": "none",
        "impact_operational": "medium",
        "impact_privacy": "none",
        "attack_feasibility": "medium",
        "remediation": "SecOC uygula.",
    }
    base.update(overrides)
    return base


def test_report_generates_valid_docx_bytes():
    profile = {"id": "p1", "name": "Test Shuttle", "architecture": "zonal"}
    findings = [_sample_finding()]
    coverage = {"total": 1, "vulnerable": 1, "not_vulnerable": 0, "vectors_tested": 1}

    docx_bytes = generate_compliance_report(profile, findings, coverage)
    assert isinstance(docx_bytes, bytes)
    assert len(docx_bytes) > 1000  # boş/bozuk dosya değil

    # python-docx ile tekrar açılabiliyor mu (dosya bütünlüğü kontrolü)
    doc = Document(BytesIO(docx_bytes))
    assert len(doc.paragraphs) > 0


def test_report_contains_key_sections():
    profile = {"id": "p1", "name": "Test Shuttle", "architecture": "zonal"}
    findings = [_sample_finding()]
    coverage = {"total": 1, "vulnerable": 1, "not_vulnerable": 0, "vectors_tested": 1}

    docx_bytes = generate_compliance_report(profile, findings, coverage)
    doc = Document(BytesIO(docx_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    assert "Yönetici Özeti" in full_text
    assert "Metodoloji" in full_text
    assert "Annex 5 Kapsam" in full_text
    assert "Test Shuttle" in full_text


def test_report_handles_empty_findings():
    profile = {"id": "p1", "name": "Bos Profil"}
    coverage = {"total": 0, "vulnerable": 0, "not_vulnerable": 0, "vectors_tested": 0}

    docx_bytes = generate_compliance_report(profile, [], coverage)
    doc = Document(BytesIO(docx_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "henüz kaydedilmiş bulgu bulunmamaktadır" in full_text


def test_report_deduplicates_repeated_findings_across_components():
    # Aynı bulgu (aynı vektör+başlık+açıklama) dört farklı bileşende tekrarlanıyor;
    # rapor bunu TEK blok olarak, bileşenleri gruplayarak göstermeli.
    findings = [
        _sample_finding(
            component_id=comp, r155_vector_id="R155-2.9",
            title="LiDAR Spoof: 2/2 senaryo başarılı",
            description="Ortak açıklama metni.", impact_safety="critical",
        )
        for comp in ["lidar_front", "lidar_rear", "camera_front", "radar_front"]
    ]
    profile = {"id": "p1", "name": "Test Shuttle"}
    coverage = {"total": 4, "vulnerable": 4, "not_vulnerable": 0, "vectors_tested": 1}

    docx_bytes = generate_compliance_report(profile, findings, coverage)
    doc = Document(BytesIO(docx_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    # Açıklama metni yalnızca BİR kez geçmeli (dört kez değil)
    assert full_text.count("Ortak açıklama metni.") == 1
    assert "lidar_front, lidar_rear, camera_front, radar_front" in full_text


def test_report_end_to_end_with_real_orchestrator(tmp_path, profile):
    db = FindingStore(db_path=str(tmp_path / "report.db"))
    orch = Orchestrator(_mock("vulnerable"), db, strict_adapter=False)
    orch.run_all(profile)

    findings = db.get_findings(vehicle_profile_id="shuttle-test")
    coverage = db.get_compliance_coverage("shuttle-test")
    profile_dict = {"id": "shuttle-test", "name": "Test Shuttle"}

    docx_bytes = generate_compliance_report(profile_dict, findings, coverage)
    doc = Document(BytesIO(docx_bytes))
    assert len(doc.tables) >= 3  # özet + bulgu + kapsam tablosu


# ── 3D Saldırı Yüzeyi ─────────────────────────────────────────────────────────

def _sample_components():
    return [
        {"id": "gateway_ecu", "label": "Gateway ECU", "category": "network",
         "position_3d": [0.2, 0.0, 0.2], "attack_surfaces": ["in-vehicle-network"],
         "r155_vectors": ["R155-2.2"]},
        {"id": "lidar_front", "label": "Ön LiDAR", "category": "sensor",
         "position_3d": [2.1, 0.0, 1.2], "attack_surfaces": ["sensor-spoofing"],
         "r155_vectors": ["R155-2.9"]},
        {"id": "obd2_port", "label": "OBD-II Portu", "category": "diagnostic",
         "position_3d": [0.8, -0.5, 0.0], "attack_surfaces": ["diagnostic"],
         "r155_vectors": ["R155-5.5"]},
        {"id": "tcu", "label": "TCU", "category": "connectivity",
         "position_3d": [0.0, 2.8, 0.5], "attack_surfaces": ["telematics"],
         "r155_vectors": ["R155-5.1"]},
    ]


def test_compute_statuses_vulnerable_takes_priority():
    components = _sample_components()
    findings = [
        {"component_id": "gateway_ecu", "status": "vulnerable"},
        {"component_id": "gateway_ecu", "status": "not_vulnerable"},
    ]
    statuses = compute_component_statuses(components, findings)
    assert statuses["gateway_ecu"] == "vulnerable"


def test_compute_statuses_clean_when_only_not_vulnerable():
    components = _sample_components()
    findings = [{"component_id": "obd2_port", "status": "not_vulnerable"}]
    statuses = compute_component_statuses(components, findings)
    assert statuses["obd2_port"] == "clean"


def test_compute_statuses_not_tested_when_no_findings():
    components = _sample_components()
    statuses = compute_component_statuses(components, [])
    assert all(s == "not_tested" for s in statuses.values())


def test_compute_statuses_ignores_inconclusive_and_error():
    components = _sample_components()
    findings = [
        {"component_id": "tcu", "status": "inconclusive"},
        {"component_id": "tcu", "status": "error"},
    ]
    statuses = compute_component_statuses(components, findings)
    assert statuses["tcu"] == "not_tested"


def test_compute_statuses_covers_all_components():
    components = _sample_components()
    statuses = compute_component_statuses(components, [])
    assert set(statuses.keys()) == {c["id"] for c in components}


def test_build_html_contains_component_data():
    components = _sample_components()
    statuses = {"gateway_ecu": "vulnerable", "lidar_front": "clean",
                "obd2_port": "not_tested", "tcu": "not_tested"}
    html = build_attack_surface_html("Test Shuttle", components, statuses)

    assert "gateway_ecu" in html
    assert "Ön LiDAR" in html
    assert "#e24b4a" in html  # vulnerable rengi
    assert "#63a922" in html  # clean rengi
    assert "#ef9f27" in html  # not_tested rengi


def test_build_html_contains_valid_importmap_json():
    import json
    import re

    components = _sample_components()
    statuses = compute_component_statuses(components, [])
    html = build_attack_surface_html("Test Shuttle", components, statuses)

    match = re.search(r'<script type="importmap">\s*(\{.*?\})\s*</script>', html, re.S)
    assert match is not None
    data = json.loads(match.group(1))  # geçerli JSON olmalı, exception atmamalı
    assert "three" in data["imports"]


def test_build_html_embeds_syntactically_valid_js():
    import shutil
    if shutil.which("node") is None:
        pytest.skip("node kurulu değil — JS sözdizimi kontrolü atlandı")

    components = _sample_components()
    statuses = compute_component_statuses(components, [])
    html = build_attack_surface_html("Test Shuttle", components, statuses)

    start = html.index('<script type="module">') + len('<script type="module">')
    end = html.index("</script>", start)
    js_code = html[start:end]

    # import ifadelerini node'un çözemeyeceği bir CDN'den çektiği için
    # sözdizimi kontrolü amacıyla nötrleştiriyoruz; geri kalan JS gövdesi
    # gerçek/değişmeden test edilir.
    js_for_check = js_code.replace(
        "import * as THREE from 'three';", "// import"
    ).replace(
        "import { OrbitControls } from 'three/addons/controls/OrbitControls.js';",
        "const THREE = {}; const OrbitControls = function(){};"
    )

    import subprocess
    import tempfile
    with tempfile.NamedTemporaryFile(suffix=".mjs", mode="w", delete=False) as f:
        f.write(js_for_check)
        path = f.name

    result = subprocess.run(["node", "--check", path], capture_output=True, text=True)
    assert result.returncode == 0, result.stderr


def test_build_html_handles_empty_components():
    html = build_attack_surface_html("Bos Profil", [], {})
    assert "gokturk-3d-root" in html


def test_compute_vector_statuses_covers_all_69():
    statuses = compute_vector_statuses([])
    assert len(statuses) == 69
    assert all(s == "not_tested" for s in statuses.values())


def test_compute_vector_statuses_vulnerable_priority():
    findings = [
        {"r155_vector_id": "R155-2.5", "status": "vulnerable"},
        {"r155_vector_id": "R155-2.5", "status": "not_vulnerable"},
    ]
    statuses = compute_vector_statuses(findings)
    assert statuses["R155-2.5"] == "vulnerable"


def test_compute_vector_statuses_clean_when_only_not_vulnerable():
    findings = [{"r155_vector_id": "R155-5.5", "status": "not_vulnerable"}]
    statuses = compute_vector_statuses(findings)
    assert statuses["R155-5.5"] == "clean"


def test_compute_vector_statuses_ignores_findings_without_vector():
    findings = [{"status": "vulnerable"}]
    statuses = compute_vector_statuses(findings)
    assert all(s == "not_tested" for s in statuses.values())


def test_build_heatmap_html_contains_all_69_cells():
    html = build_heatmap_html([])
    assert html.count('gk-heat-cell"') == 69


def test_build_heatmap_html_reflects_findings():
    findings = [{"r155_vector_id": "R155-2.5", "status": "vulnerable"}]
    html = build_heatmap_html(findings)
    assert "var(--red)" in html
    assert "R155-2.5" in html


def test_build_heatmap_html_escapes_tooltip_content():
    html = build_heatmap_html([])
    import re
    titles = re.findall(r'title="([^"]*)"', html)
    assert len(titles) == 69


def test_build_heatmap_html_uses_theme_css_vars_not_hardcoded_colors():
    html = build_heatmap_html([])
    assert "var(--red)" in html
    assert "var(--green)" in html
    assert "var(--text-muted)" in html


# ── Otonom Servis Aracı v1 Profili ─────────────────────────────────────────────

def _load_autonomous_shuttle():
    import yaml
    with open("profiles/autonomous_shuttle_v1.yaml", encoding="utf-8") as f:
        return yaml.safe_load(f.read())


def test_autonomous_shuttle_profile_parses():
    prof = _load_autonomous_shuttle()
    assert prof["id"] == "autonomous-shuttle-v1"
    assert "components" in prof


def test_autonomous_shuttle_has_at_least_10_components():
    prof = _load_autonomous_shuttle()
    assert len(prof["components"]) >= 10


def test_autonomous_shuttle_components_have_required_fields():
    prof = _load_autonomous_shuttle()
    required = {"id", "label", "category", "position_3d", "attack_surfaces",
                "r155_vectors", "networks", "test_status"}
    for comp in prof["components"]:
        missing = required - set(comp.keys())
        assert not missing, f"{comp.get('id')} eksik alanlar: {missing}"


def test_autonomous_shuttle_component_ids_unique():
    prof = _load_autonomous_shuttle()
    ids = [c["id"] for c in prof["components"]]
    assert len(ids) == len(set(ids))


def test_autonomous_shuttle_all_r155_vectors_exist_in_taxonomy():
    prof = _load_autonomous_shuttle()
    tax = load_taxonomy()
    valid_ids = {v["id"] for cat in tax["categories"] for v in cat["vectors"]}
    for comp in prof["components"]:
        for vid in comp["r155_vectors"]:
            assert vid in valid_ids, f"{comp['id']}: gecersiz vektor {vid}"


def test_autonomous_shuttle_runs_through_orchestrator(tmp_path):
    prof = _load_autonomous_shuttle()
    db = FindingStore(db_path=str(tmp_path / "shuttle_v1.db"))
    orch = Orchestrator(_mock("vulnerable"), db, strict_adapter=False)
    findings = orch.run_all(prof)
    assert len(findings) > 0
    assert all(isinstance(f, Finding) for f in findings)


def test_autonomous_shuttle_compatible_with_3d_map_and_heatmap(tmp_path):
    prof = _load_autonomous_shuttle()
    db = FindingStore(db_path=str(tmp_path / "shuttle_v1.db"))
    orch = Orchestrator(_mock("vulnerable"), db, strict_adapter=False)
    orch.run_all(prof)

    findings = db.get_findings()
    statuses = compute_component_statuses(prof["components"], findings)
    html3d = build_attack_surface_html(prof["name"], prof["components"], statuses)
    assert "gokturk-3d-root" in html3d
    for comp in prof["components"]:
        assert comp["id"] in html3d

    heat = build_heatmap_html(findings)
    assert heat.count('gk-heat-cell"') == 69


# ── TARA Belge Üreticisi ───────────────────────────────────────────────────────

def test_tara_generator_produces_document():
    from scripts.generate_tara import generate
    doc = generate()
    assert "Tehdit Analizi ve Risk Değerlendirmesi" in doc
    assert "ISO/SAE 21434" in doc
    assert "autonomous-shuttle-v1" in doc


def test_tara_lists_all_10_components():
    from scripts.generate_tara import generate
    import yaml
    with open("profiles/autonomous_shuttle_v1.yaml", encoding="utf-8") as f:
        prof = yaml.safe_load(f.read())
    doc = generate()
    for comp in prof["components"]:
        assert comp["id"] in doc


def test_tara_safety_floor_no_critical_safety_below_high():
    # Güvenlik etkisi critical/high olan hiçbir tehdit 'Orta' veya altı olmamalı
    from scripts.generate_tara import generate
    doc = generate()
    for line in doc.splitlines():
        if line.startswith("| **Orta**") or line.startswith("| **Düşük**"):
            assert "G:critical" not in line


# ── Firmware Integrity Plugin ──────────────────────────────────────────────────

def test_firmware_integrity_matrix():
    vuln = FirmwareIntegrityPlugin(_mock("vulnerable")).run({"id": "hpc_compute"})
    assert isinstance(vuln, list) and all(f.status == "vulnerable" for f in vuln)
    assert FirmwareIntegrityPlugin(_mock("secure")).run({"id": "hpc_compute"}).status == "not_vulnerable"
    assert FirmwareIntegrityPlugin(_mock("empty")).run({"id": "hpc_compute"}).status == "not_vulnerable"


def test_firmware_integrity_vulnerable_returns_three_distinct_vectors():
    findings = FirmwareIntegrityPlugin(_mock("vulnerable")).run({"id": "hpc_compute"})
    vectors = sorted(f.r155_vector_id for f in findings)
    assert vectors == ["R155-6.1", "R155-6.13", "R155-6.4"]
    assert all(f.r155_category == 6 for f in findings)
    replace_finding = next(f for f in findings if f.r155_vector_id == "R155-6.1")
    assert replace_finding.impact_safety == "critical"
    boot_finding = next(f for f in findings if f.r155_vector_id == "R155-6.13")
    assert boot_finding.impact_safety == "critical"


def test_firmware_integrity_lists_all_three_scenarios():
    findings = FirmwareIntegrityPlugin(_mock("vulnerable")).run({"id": "hpc_compute"})
    assert len(findings) == 3


def test_firmware_integrity_secure_reports_all_protected():
    f = FirmwareIntegrityPlugin(_mock("secure")).run({"id": "hpc_compute"})
    assert f.status == "not_vulnerable"
    assert "korumalar" in f.title.lower() or "aktif" in f.title.lower()


def test_firmware_integrity_in_discovery():
    from core.orchestrator import Orchestrator
    from adapters.mock_adapter import MockAdapter
    orch = Orchestrator(MockAdapter({"mode": "vulnerable"}), None, strict_adapter=False)
    ids = {c.module_id for c in orch.discover_plugin_classes()}
    assert "firmware-integrity" in ids


# ── Mock: remote_telematics_exploit_probe davranışı ────────────────────────────

def test_mock_remote_telematics_exploit_probe_behaviour():
    assert _mock("vulnerable").remote_telematics_exploit_probe("telematics_module") is True
    assert _mock("secure").remote_telematics_exploit_probe("telematics_module") is False
    assert _mock("empty").remote_telematics_exploit_probe("telematics_module") is False


# ── Remote Telematics Exploit Plugin ────────────────────────────────────────────

def test_remote_telematics_exploit_matrix():
    assert RemoteTelematicsExploitPlugin(_mock("vulnerable")).run({"id": "telematics_module"}).status == "vulnerable"
    assert RemoteTelematicsExploitPlugin(_mock("secure")).run({"id": "telematics_module"}).status == "not_vulnerable"
    assert RemoteTelematicsExploitPlugin(_mock("empty")).run({"id": "telematics_module"}).status == "not_vulnerable"


def test_remote_telematics_exploit_carries_taxonomy():
    f = RemoteTelematicsExploitPlugin(_mock("vulnerable")).run({"id": "telematics_module"})
    assert f.r155_vector_id == "R155-5.13"
    assert f.r155_category == 5
    assert f.impact_safety == "critical"
    assert f.is_vulnerable()
    assert f.is_critical_safety()


def test_remote_telematics_exploit_inconclusive_when_adapter_unsupported():
    from adapters.base_adapter import BaseAdapter

    class _BareAdapter(BaseAdapter):
        adapter_type = "bare"

        def connect(self):
            self._connected = True
            return True

        def disconnect(self):
            self._connected = False

        def is_connected(self):
            return self._connected

    bare = _BareAdapter({})
    bare.connect()
    f = RemoteTelematicsExploitPlugin(bare).run({"id": "telematics_module"})
    assert f.status == "inconclusive"


def test_remote_telematics_exploit_in_discovery():
    from core.orchestrator import Orchestrator
    from adapters.mock_adapter import MockAdapter
    orch = Orchestrator(MockAdapter({"mode": "vulnerable"}), None, strict_adapter=False)
    ids = {c.module_id for c in orch.discover_plugin_classes()}
    assert "remote-telematics-exploit" in ids


# ── Mock: can_dos_probe davranışı ───────────────────────────────────────────────

def test_mock_can_dos_probe_behaviour():
    for tech in ("high_priority_flood", "error_frame_attack"):
        assert _mock("vulnerable").can_dos_probe("gateway_ecu", tech)["succeeded"] is True
        assert _mock("secure").can_dos_probe("gateway_ecu", tech)["succeeded"] is False
        assert _mock("empty").can_dos_probe("gateway_ecu", tech)["succeeded"] is False


# ── CAN DoS Plugin ───────────────────────────────────────────────────────────────

def test_can_dos_matrix():
    assert CANDosPlugin(_mock("vulnerable")).run({"id": "gateway_ecu"}).status == "vulnerable"
    assert CANDosPlugin(_mock("secure")).run({"id": "gateway_ecu"}).status == "not_vulnerable"
    assert CANDosPlugin(_mock("empty")).run({"id": "gateway_ecu"}).status == "not_vulnerable"


def test_can_dos_carries_taxonomy():
    f = CANDosPlugin(_mock("vulnerable")).run({"id": "gateway_ecu"})
    assert f.r155_vector_id == "R155-2.4"
    assert f.r155_category == 2
    assert f.impact_safety == "high"
    assert f.is_vulnerable()


def test_can_dos_error_frame_technique_config():
    plugin = CANDosPlugin(_mock("vulnerable"), config={"dos_technique": "error_frame_attack"})
    f = plugin.run({"id": "gateway_ecu"})
    assert f.status == "vulnerable"
    assert "bus-off" in f.title.lower() or "hata çerçevesi" in f.title.lower()


def test_can_dos_in_discovery():
    from core.orchestrator import Orchestrator
    from adapters.mock_adapter import MockAdapter
    orch = Orchestrator(MockAdapter({"mode": "vulnerable"}), None, strict_adapter=False)
    ids = {c.module_id for c in orch.discover_plugin_classes()}
    assert "can-dos" in ids


# ── Mock: ivi_pivot_probe davranışı ─────────────────────────────────────────────

def test_mock_ivi_pivot_probe_behaviour():
    assert _mock("vulnerable").ivi_pivot_probe("gateway_ecu") is True
    assert _mock("secure").ivi_pivot_probe("gateway_ecu") is False
    assert _mock("empty").ivi_pivot_probe("gateway_ecu") is False


# ── IVI Pivot Plugin ─────────────────────────────────────────────────────────────

def test_ivi_pivot_matrix():
    assert IVIPivotPlugin(_mock("vulnerable")).run({"id": "gateway_ecu"}).status == "vulnerable"
    assert IVIPivotPlugin(_mock("secure")).run({"id": "gateway_ecu"}).status == "not_vulnerable"
    assert IVIPivotPlugin(_mock("empty")).run({"id": "gateway_ecu"}).status == "not_vulnerable"


def test_ivi_pivot_carries_taxonomy():
    f = IVIPivotPlugin(_mock("vulnerable")).run({"id": "gateway_ecu"})
    assert f.r155_vector_id == "R155-5.4"
    assert f.r155_category == 5
    assert f.impact_safety == "high"
    assert f.is_vulnerable()


def test_ivi_pivot_in_discovery():
    from core.orchestrator import Orchestrator
    from adapters.mock_adapter import MockAdapter
    orch = Orchestrator(MockAdapter({"mode": "vulnerable"}), None, strict_adapter=False)
    ids = {c.module_id for c in orch.discover_plugin_classes()}
    assert "ivi-pivot" in ids
# ── Mock: telematics_channel_probe davranışı ────────────────────────────────────

def test_mock_telematics_channel_probe_behaviour():
    assert _mock("vulnerable").telematics_channel_probe("telematics_module") is True
    assert _mock("secure").telematics_channel_probe("telematics_module") is False
    assert _mock("empty").telematics_channel_probe("telematics_module") is False


# ── Telematics Channel Plugin ────────────────────────────────────────────────────

def test_telematics_channel_matrix():
    assert TelematicsChannelPlugin(_mock("vulnerable")).run({"id": "telematics_module"}).status == "vulnerable"
    assert TelematicsChannelPlugin(_mock("secure")).run({"id": "telematics_module"}).status == "not_vulnerable"
    assert TelematicsChannelPlugin(_mock("empty")).run({"id": "telematics_module"}).status == "not_vulnerable"


def test_telematics_channel_carries_taxonomy():
    f = TelematicsChannelPlugin(_mock("vulnerable")).run({"id": "telematics_module"})
    assert f.r155_vector_id == "R155-5.1"
    assert f.r155_category == 5
    assert f.impact_privacy == "high"
    assert f.is_vulnerable()


def test_telematics_channel_in_discovery():
    from core.orchestrator import Orchestrator
    from adapters.mock_adapter import MockAdapter
    orch = Orchestrator(MockAdapter({"mode": "vulnerable"}), None, strict_adapter=False)
    ids = {c.module_id for c in orch.discover_plugin_classes()}
    assert "telematics-channel" in ids
